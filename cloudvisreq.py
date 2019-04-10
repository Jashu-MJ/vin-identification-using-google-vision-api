from base64 import b64encode
from os import makedirs
from os.path import join, basename
from PIL import Image
from docx import Document
from docx.shared import Inches
#from sys import argv

import json
import requests
import re
import glob


ENDPOINT_URL = 'https://vision.googleapis.com/v1/images:annotate'
RESULTS_DIR = 'jsons'
makedirs(RESULTS_DIR, exist_ok=True)

def resize():
    for filename in glob.glob(r'C:\Users\Jaswanth Marri\My Documents\LiClipse Workspace\First.py\vin_number_recognition\*JPG'): #assuming gif
        foo = Image.open(filename)
        w,h=foo.size
        if(w == 4608):
            foo = Image.open(filename)
            print('dimensions of the image before',foo.size)
        # I downsize the image with an ANTIALIAS filter (gives the highest quality)
            foo = foo.resize((640,480),Image.ANTIALIAS)
            print('dimension of the image after',foo.size)
            foo.save(filename,optimize=True,quality=100)

def make_image_data_list(image_filenames):
    """
    image_filenames is a list of filename strings
    Returns a list of dicts formatted as the Vision API
        needs them to be
    """
    img_requests = []
    for imgname in image_filenames:
        with open(imgname, 'rb') as f:
            ctxt = b64encode(f.read()).decode()
            img_requests.append({
                    'image': {'content': ctxt},
                    'features': [{
                        'type': 'TEXT_DETECTION',
                        'maxResults': 1
                    }]
            })
    return img_requests

def make_image_data(image_filenames):
    """Returns the image data lists as bytes"""
    imgdict = make_image_data_list(image_filenames)
    return json.dumps({"requests": imgdict }).encode()


def request_ocr(api_key, image_filenames):
    response = requests.post(ENDPOINT_URL,
                             data=make_image_data(image_filenames),
                             params={'key': api_key},
                             headers={'Content-Type': 'application/json'})
    return response

def divide_chunks(l, n): 
    
    # looping till length l 
    for i in range(0, len(l), n): 
        yield l[i:i + n] 

if __name__ == '__main__':
    # here enter the google cloud vision api key so that the program will get connected with the API
    api_key = "Enter api key"
    #list of all images that has to be processed initially empty give the location of folder with images in for loop and resize function
    image_filenames=[]
    #file1 = open("myfile.txt","w")
    #if the image file is large then it will take more time for uploading so optimising the image and resizing the image to reduce the size
    # also an image size should not exceed more than 20 mb
    ###################################***********************
    #enter the location of the images in the fuction definition
    #this function will resize if the image is 4608 x 3456 to 640*480
    resize()
    ######################**********************************
    #print(type(image_filenames))
    document = Document()

    document.add_heading('Results', 0)
    
    #specify the location of the folder with images
    for filename in glob.glob(r'C:\Users\Jaswanth Marri\My Documents\LiClipse Workspace\First.py\vin_number_recognition\*JPG'): #assuming jpg
    #for filename in glob.glob(r'https://drive.google.com/open?id=174JBzNa1A_OYCD-WrmJ8tiiI7hklkk_p/*JPG'): #assuming jpg

            image_filenames.append(filename)
            
    #n is the number of images that will be read in one itteration as the limit for vision api is 16 images at once
    # one JSON request object size should not exceed more than 10 mb so we divide it into batches
    n=16
    #dividing the list into batches of required size
    x = list(divide_chunks(image_filenames, n))
    #print(image_filenames)
    #print(type(image_filenames))
    #print(type(*image_filenames))
    #print(type(api_key))
    #print(x)
    s="size per batch:"+str(n)
    print(s)
    document.add_paragraph(s,style='List Bullet')
    q='the no of batches:'+str(len(x))
    print(q)
    document.add_paragraph(q,style='List Bullet')

    
    if not api_key or not image_filenames:
        print("""
            Please supply an api key, then one or more image filenames

            """)
    else:
        for i in range(len(x)):
            it="itterating loop: "+str(i)
            print(it)
            document.add_paragraph(it,style='List Bullet')
            ni='the no of images in batch :'+str(len(x[i]))
            print(ni)
            document.add_paragraph(ni,style='List Bullet')
            response = request_ocr(api_key,x[i])
            if response.status_code != 200 or response.json().get('error'):
                print(response.text)
            else:
                for idx, resp in enumerate(response.json()['responses']):
                    # save to JSON file
                    #print(idx)
                    #print(type(idx))
                    imgname = x[i][idx]
                    jpath = join(RESULTS_DIR, basename(imgname) + '.json')
                    with open(jpath, 'w') as f:
                        datatxt = json.dumps(resp, indent=2)
                        print("Wrote", len(datatxt), "bytes to", jpath)
                        document.add_paragraph("wrote"+str(len(datatxt))+"bytes to"+str(jpath),style='List Bullet')
                        f.write(datatxt)
    
                    # print the plaintext to screen for convenience
                    print("---------------------------------------------")
                    
                    t = resp['textAnnotations'][0]
                    '''
                    print("    Bounding Polygon:")
                    print(t['boundingPoly'])
                    print("    Text:")
                    print(t['description'])
                    '''
                    print(t['description'])
                    document.add_paragraph("---------------------------------------------\n",style='List Bullet')
                    document.add_paragraph("  TEXT :", style='List Bullet')
                    document.add_paragraph(t['description'], style='List Number')
                    #file1.write('text:','t['description'])

                    #print(type(t['description']))
                    print('the vin number is')
                    a=re.findall('[A-Z]{3}[A-Z0-9]{10}[0-9]{4}',t['description'])
                    print(a)
                    document.add_paragraph('the vin number is ', style='List Bullet')
                    document.add_paragraph(a, style='List Bullet')
                    document.save('output.docx')

                    #file1.write('the vin number is')
                    #file1.writelines(re.findall('[A-Z]{3}[A-Z0-9]{10}[0-9]{4}',t['description']))

                    #print((re.findall('[A-Z]{3}[A-Z0-9]{10}[0-9]{4}',t['description'])).isna())
    #file1.close()
                